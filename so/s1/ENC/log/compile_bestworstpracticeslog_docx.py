#!/usr/bin/env python3
"""
Compiles bestworstpracticeslog.md into an MLA 9th Edition formatted DOCX / DOC document.
Follows the specifications in summaries/GEMINI.md.
"""

import os
import shutil
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def build_docx():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    output_docx_path = os.path.join(base_dir, "bestworstpracticeslog.docx")
    output_doc_path = os.path.join(base_dir, "bestworstpracticeslog.doc")

    doc = docx.Document()

    # Page setup (US Letter, 1-inch margins)
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.5)

        # Header: Lopez 1 (MLA running head)
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.paragraph_format.space_before = Pt(0)
        hp.paragraph_format.space_after = Pt(0)
        hp.paragraph_format.line_spacing = 1.0
        hrun = hp.add_run("Lopez ")
        hrun.font.name = "Times New Roman"
        hrun.font.size = Pt(12)

        # Dynamic page number field
        fldSimple = OxmlElement('w:fldSimple')
        fldSimple.set(qn('w:instr'), 'PAGE')
        hp._p.append(fldSimple)

    # Set Normal style font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    def add_mla_paragraph(text="", align=WD_ALIGN_PARAGRAPH.LEFT, indent=0.0, bold=False, italic=False):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.first_line_indent = Inches(indent)
        if text:
            r = p.add_run(text)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
            r.bold = bold
            r.italic = italic
        return p

    # Student Identification Block
    add_mla_paragraph("Owen Lopez")
    add_mla_paragraph("Professor Faulconer")
    add_mla_paragraph("ENC 3250")
    add_mla_paragraph("29 August 2026")

    # Document Title
    add_mla_paragraph("Best and Worst Practices Log", align=WD_ALIGN_PARAGRAPH.CENTER)

    # Best Practices Section
    add_mla_paragraph("Best Practices", bold=True)
    p_bp = add_mla_paragraph(indent=0.5)
    r1 = p_bp.add_run("Vocabulary: ")
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(12)
    r1.italic = True
    r2 = p_bp.add_run(
        "In almost every serious sentence I put down I will review for grammar, "
        "punctuation, & specifically for vocabulary. I ensure that every verb, noun, or adverb "
        "is specifically chosen to reflect the context of my current piece and always aims for a high "
        "level of complexity/detail. It is very rare for me to utilize a \"simple\" word or phrase "
        "when I am writing, as I really do spend a considerable amount of time choosing the proper, "
        "advanced, term for the situation."
    )
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(12)

    # Worst Practices Section
    add_mla_paragraph("Worst Practices", bold=True)
    p_wp = add_mla_paragraph(indent=0.5)
    r3 = p_wp.add_run("Run-On Sentences: ")
    r3.font.name = "Times New Roman"
    r3.font.size = Pt(12)
    r3.italic = True
    r4 = p_wp.add_run(
        "Run-On sentences are when an author places too many clauses into a single sentence, "
        "through commas or repetitive uses of conjunctions, instead of utilizing multiple "
        "related sentences. I often find myself writing Run-On sentences when I want to express "
        "a complex though. I have to slow down, evaluate my last sentence, and split it into "
        "multiple complete sentences."
    )
    r4.font.name = "Times New Roman"
    r4.font.size = Pt(12)

    # Save documents
    doc.save(output_docx_path)
    shutil.copyfile(output_docx_path, output_doc_path)
    print(f"Generated: {output_docx_path}")
    print(f"Generated: {output_doc_path}")

if __name__ == "__main__":
    build_docx()
